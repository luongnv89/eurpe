"""Tests for :class:`eurpe.export.docx.DocxCitationRenderer`.

Pin the issue #17 acceptance criteria at the renderer level:

* **AC #2** — :meth:`render` returns a valid DOCX byte stream paired
  with a shadow Markdown mirror.
* **AC #3** — citations and source-status labels survive into both
  the DOCX paragraph tree (so the user sees them when they open the
  file in Word/LibreOffice) and the shadow string (so the citation
  audit can apply PRD §22 checks against a single textual surface).

These tests exercise the renderer directly (no service, no audit).
The service-layer tests in ``test_export_service.py`` cover the
audit + service-injection seams; here we focus on the docx structural
guarantees so a regression in either path is caught at the right
layer.

Why round-trip via ``docx.Document(BytesIO(...))``
--------------------------------------------------
python-docx's saved files are not byte-deterministic across versions
(``app.xml`` / ``core.xml`` carry the package version), so a raw
``bytes ==`` assertion would be brittle. Loading the bytes back and
asserting on paragraph/table contents proves the user-visible
structure without depending on the XML representation.
"""

from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document

from eurpe.export.docx import DocxCitationRenderer
from eurpe.generation.models import CitationRef, GenerationDraft, GenerationRequest
from eurpe.generation.render import STATUS_BADGE, STATUS_LABEL, MarkdownCitationRenderer
from eurpe.schema import Programme, SectionType, SourceStatus

# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------


def _make_request(
    section_type: SectionType = SectionType.METHODOLOGY,
) -> GenerationRequest:
    return GenerationRequest(
        section_type=section_type,
        user_intent="describe the methodology",
        target_programme=Programme.HORIZON_EUROPE,
    )


def _make_citation(
    citation_id: int,
    *,
    source_status: SourceStatus = SourceStatus.FUNDED,
    page: int | None = 12,
    section_heading: str | None = "1.2 Methodology",
    proposal_title: str | None = "SAMPLE PROPOSAL",
) -> CitationRef:
    return CitationRef(
        citation_id=citation_id,
        source_status=source_status,
        programme=Programme.HORIZON_EUROPE,
        call_id="HORIZON-CL5-2024-D3-02",
        proposal_title=proposal_title,
        section_heading=section_heading,
        page=page,
        chunk_id=f"chunk-{citation_id}",
        snippet=f"snippet {citation_id}",
    )


def _make_draft(
    *,
    citations: list[CitationRef] | None = None,
    text: str | None = None,
    section_type: SectionType = SectionType.METHODOLOGY,
) -> GenerationDraft:
    citations = citations if citations is not None else [_make_citation(1)]
    if text is None:
        if citations:
            markers = " ".join(f"[{c.citation_id}]" for c in citations)
            text = f"Methodology overview with references {markers}."
        else:
            text = "Body without citations."
    return GenerationDraft(
        section_type=section_type,
        text=text,
        citations=citations,
        prompt_used="prompt-for-test",
        model="deterministic-stub-v1",
        request=_make_request(section_type=section_type),
    )


def _load_paragraphs(docx_bytes: bytes) -> list[str]:
    document = Document(BytesIO(docx_bytes))
    return [p.text for p in document.paragraphs]


def _load_table_rows(docx_bytes: bytes) -> list[list[str]]:
    document = Document(BytesIO(docx_bytes))
    if not document.tables:
        return []
    return [[cell.text for cell in row.cells] for row in document.tables[0].rows]


# ---------------------------------------------------------------------------
# AC #2 — DOCX renders to a valid binary stream + shadow Markdown
# ---------------------------------------------------------------------------


def test_render_returns_bytes_and_shadow_markdown() -> None:
    """The renderer returns a ``(bytes, str)`` tuple, both non-empty."""

    renderer = DocxCitationRenderer()
    docx_bytes, shadow_md = renderer.render(_make_draft())

    assert isinstance(docx_bytes, bytes)
    assert isinstance(shadow_md, str)
    assert len(docx_bytes) > 0
    assert len(shadow_md) > 0
    # DOCX is a ZIP container — the magic number is "PK\x03\x04".
    assert docx_bytes.startswith(b"PK\x03\x04")


def test_render_bytes_round_trip_through_docx_loader() -> None:
    """The output opens cleanly via ``docx.Document`` (real Word would too)."""

    renderer = DocxCitationRenderer()
    docx_bytes, _ = renderer.render(_make_draft())

    # If python-docx can parse it, Word and LibreOffice can too — the
    # file format is shared. A raised exception here would mean the
    # renderer wrote a malformed package.
    document = Document(BytesIO(docx_bytes))
    assert document is not None


def test_render_emits_section_heading_as_paragraph() -> None:
    """The section heading is the first textual paragraph in the docx."""

    renderer = DocxCitationRenderer()
    docx_bytes, _ = renderer.render(_make_draft(section_type=SectionType.METHODOLOGY))

    paragraphs = _load_paragraphs(docx_bytes)
    assert paragraphs[0] == "Methodology"


def test_render_emits_body_text_with_inline_markers_preserved() -> None:
    """Body paragraph(s) keep the inline ``[N]`` markers verbatim."""

    renderer = DocxCitationRenderer()
    docx_bytes, _ = renderer.render(
        _make_draft(
            text="Body talks about reference [1] and reference [2].",
            citations=[_make_citation(1), _make_citation(2)],
        ),
    )

    paragraphs = _load_paragraphs(docx_bytes)
    body_paragraphs = [p for p in paragraphs if "[1]" in p or "[2]" in p]
    # The body paragraph survived (we expect at least one; the bullet
    # paragraphs in the Notes block also contain markers but the body
    # paragraph is the one with the prose).
    assert any("Body talks about reference [1] and reference [2]." == p for p in body_paragraphs)


def test_render_section_imipact_pathway_title_cases() -> None:
    """``IMPACT_PATHWAY`` becomes ``"Impact Pathway"`` (matches MD renderer)."""

    renderer = DocxCitationRenderer()
    docx_bytes, _ = renderer.render(_make_draft(section_type=SectionType.IMPACT_PATHWAY))
    paragraphs = _load_paragraphs(docx_bytes)
    assert paragraphs[0] == "Impact Pathway"


# ---------------------------------------------------------------------------
# AC #3 — citation + source-status preservation
# ---------------------------------------------------------------------------


def test_render_references_table_has_seven_columns_per_row() -> None:
    """The References table mirrors the Markdown 7-column layout."""

    renderer = DocxCitationRenderer()
    docx_bytes, _ = renderer.render(_make_draft())

    rows = _load_table_rows(docx_bytes)
    # 1 header row + 1 citation row = 2.
    assert len(rows) == 2
    header, first_row = rows[0], rows[1]
    assert header == ["#", "Status", "Programme", "Call", "Section", "Page", "Source"]
    assert len(first_row) == 7
    # Status column carries the same label string the Markdown table
    # emits — this is the AC #3 contract.
    assert first_row[1] == STATUS_LABEL[SourceStatus.FUNDED]


@pytest.mark.parametrize(
    "status",
    [SourceStatus.FUNDED, SourceStatus.REJECTED, SourceStatus.ESR_NOTE],
)
def test_render_preserves_each_source_status_label(status: SourceStatus) -> None:
    """Every supported status emits its label in the docx table.

    AC #3 specifically calls out funded/rejected/ESR — pin each one
    so a future refactor that drops a status string is caught here
    rather than at a user-visible review.
    """

    renderer = DocxCitationRenderer()
    docx_bytes, _ = renderer.render(
        _make_draft(
            citations=[_make_citation(1, source_status=status)],
            text="Body talks about reference [1].",
        ),
    )

    rows = _load_table_rows(docx_bytes)
    assert rows[1][1] == STATUS_LABEL[status]


@pytest.mark.parametrize(
    "status",
    [SourceStatus.FUNDED, SourceStatus.REJECTED, SourceStatus.ESR_NOTE],
)
def test_render_notes_bullet_carries_status_badge(status: SourceStatus) -> None:
    """The Notes paragraph for each citation includes the status badge.

    Mirrors the Markdown notes block — the glyph + label travels into
    the docx so a Word reader sees the same vocabulary as the CLI.
    """

    renderer = DocxCitationRenderer()
    docx_bytes, _ = renderer.render(
        _make_draft(
            citations=[_make_citation(1, source_status=status)],
            text="Body talks about reference [1].",
        ),
    )

    paragraphs = _load_paragraphs(docx_bytes)
    badge = STATUS_BADGE[status]
    assert any(
        badge in p for p in paragraphs
    ), f"Expected status badge {badge!r} in DOCX paragraphs, got: {paragraphs}"


@pytest.mark.parametrize(
    "status,caveat_fragment",
    [
        (SourceStatus.REJECTED, "Cautionary lesson"),
        (SourceStatus.ESR_NOTE, "Reviewer commentary"),
    ],
)
def test_render_emits_caveat_paragraph_for_non_funded_status(
    status: SourceStatus, caveat_fragment: str
) -> None:
    """Non-funded statuses emit an italic caveat paragraph in the docx.

    The Markdown renderer attaches a ``_…_`` italic line; the DOCX
    renderer strips the underscores (italic is a run flag in docx)
    and surfaces the same prose. The caveat fragment is the load-
    bearing word a reviewer would scan for.
    """

    renderer = DocxCitationRenderer()
    docx_bytes, _ = renderer.render(
        _make_draft(
            citations=[_make_citation(1, source_status=status)],
            text="Body talks about reference [1].",
        ),
    )

    paragraphs = _load_paragraphs(docx_bytes)
    assert any(caveat_fragment in p for p in paragraphs)


def test_render_funded_status_does_not_emit_caveat_paragraph() -> None:
    """Funded citations skip the caveat paragraph (no underscore line).

    The Markdown renderer emits an empty caveat string for FUNDED,
    and the DOCX renderer mirrors that. A funded draft should not
    carry a "cautionary lesson" paragraph.
    """

    renderer = DocxCitationRenderer()
    docx_bytes, _ = renderer.render(_make_draft())

    paragraphs = _load_paragraphs(docx_bytes)
    assert not any("Cautionary lesson" in p for p in paragraphs)
    assert not any("Reviewer commentary" in p for p in paragraphs)


# ---------------------------------------------------------------------------
# Shadow Markdown — single source of truth for the audit
# ---------------------------------------------------------------------------


def test_shadow_markdown_matches_standalone_markdown_renderer() -> None:
    """The shadow string is byte-equal to MarkdownCitationRenderer's output.

    Documents the contract the export service relies on: feeding the
    shadow through :meth:`CitationAudit.audit_rendered` runs the same
    PRD §22 checks the Markdown branch sees. A drift between the two
    would invalidate that contract silently.
    """

    draft = _make_draft()
    renderer = DocxCitationRenderer()
    _, shadow_md = renderer.render(draft)

    md_renderer = MarkdownCitationRenderer()
    assert shadow_md == md_renderer.render(draft)


def test_shadow_markdown_contains_status_badge_for_each_citation() -> None:
    """Every cited status's badge appears in the shadow string.

    The export service feeds this string to ``audit_rendered``; if a
    badge were missing the audit would refuse the export. This test
    pins the property up-front so a regression here surfaces a clean
    diff before it cascades into an audit failure.
    """

    renderer = DocxCitationRenderer()
    citations = [
        _make_citation(1, source_status=SourceStatus.FUNDED),
        _make_citation(2, source_status=SourceStatus.REJECTED),
        _make_citation(3, source_status=SourceStatus.ESR_NOTE),
    ]
    _, shadow_md = renderer.render(
        _make_draft(citations=citations, text="Body talks about [1] [2] [3]."),
    )

    for citation in citations:
        assert STATUS_BADGE[citation.source_status] in shadow_md


# ---------------------------------------------------------------------------
# Edge cases
# ---------------------------------------------------------------------------


def test_render_handles_empty_citation_list() -> None:
    """A draft with zero citations renders without raising.

    The Markdown renderer emits ``_No citations._`` for the empty
    case; the DOCX renderer emits an italic ``"No citations."``
    paragraph. The audit treats the empty case as legal so the export
    service hands back a valid result.
    """

    renderer = DocxCitationRenderer()
    docx_bytes, shadow_md = renderer.render(
        _make_draft(citations=[], text="Body without any references."),
    )

    # Shadow string keeps the underscore form so it matches the
    # Markdown renderer byte-for-byte.
    assert "_No citations._" in shadow_md
    # DOCX paragraphs carry the prose without the underscores.
    paragraphs = _load_paragraphs(docx_bytes)
    assert "No citations." in paragraphs
    # No table is emitted when there are no citations.
    rows = _load_table_rows(docx_bytes)
    assert rows == []


def test_render_missing_page_falls_back_to_placeholder() -> None:
    """A citation with ``page=None`` renders ``"n/a"`` in the table cell.

    Mirrors the Markdown table — a bare ``None`` would otherwise leak
    into the document. The Notes paragraph uses the ``p. n/a`` form
    that matches the Markdown convention.
    """

    renderer = DocxCitationRenderer()
    docx_bytes, _ = renderer.render(
        _make_draft(citations=[_make_citation(1, page=None)], text="Body about [1]."),
    )

    rows = _load_table_rows(docx_bytes)
    assert rows[1][5] == "n/a"
    paragraphs = _load_paragraphs(docx_bytes)
    assert any("p. n/a" in p for p in paragraphs)


def test_render_missing_section_and_title_render_placeholders() -> None:
    """Missing section heading / proposal title render their placeholders."""

    renderer = DocxCitationRenderer()
    docx_bytes, _ = renderer.render(
        _make_draft(
            citations=[_make_citation(1, section_heading=None, proposal_title=None)],
            text="Body about [1].",
        ),
    )

    rows = _load_table_rows(docx_bytes)
    assert rows[1][4] == "n/a"
    assert rows[1][6] == "untitled"


def test_render_multiple_citations_emit_one_table_row_per_citation() -> None:
    """N citations → N+1 rows (header + N) in the references table."""

    renderer = DocxCitationRenderer()
    citations = [_make_citation(i) for i in range(1, 4)]
    docx_bytes, _ = renderer.render(
        _make_draft(citations=citations, text="[1] [2] [3]"),
    )

    rows = _load_table_rows(docx_bytes)
    assert len(rows) == 1 + len(citations)


def test_render_body_with_blank_separated_paragraphs_splits_correctly() -> None:
    """Double-newline-separated body becomes multiple DOCX paragraphs.

    A coordinator who pastes the section into the Word template
    expects paragraph breaks to round-trip — collapsing them into a
    single paragraph would change the visible layout.
    """

    renderer = DocxCitationRenderer()
    body = "First paragraph with [1].\n\nSecond paragraph.\n\nThird paragraph."
    docx_bytes, _ = renderer.render(_make_draft(text=body))

    paragraphs = _load_paragraphs(docx_bytes)
    assert "First paragraph with [1]." in paragraphs
    assert "Second paragraph." in paragraphs
    assert "Third paragraph." in paragraphs


def test_render_stateless_across_calls() -> None:
    """Two consecutive renders on the same draft produce equal-length output.

    The renderer is stateless; the second render must not accumulate
    paragraphs from the first. Equal lengths are a cheap structural
    check (raw bytes are not deterministic across python-docx versions,
    but length is stable within one process run).
    """

    renderer = DocxCitationRenderer()
    draft = _make_draft()
    first, _ = renderer.render(draft)
    second, _ = renderer.render(draft)
    assert len(first) == len(second)
