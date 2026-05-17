"""DOCX rendering for generated drafts with visible source-status labels.

The DOCX renderer is the binary counterpart to
:class:`~eurpe.generation.render.MarkdownCitationRenderer`. Both produce
output a coordinator can hand off to a co-author who is preparing the
official EU proposal template — Markdown for git-friendly editing,
DOCX so users can paste a polished section directly into the Word /
LibreOffice template the EU portal expects.

What this module guarantees (issue #17 acceptance criteria)
-----------------------------------------------------------
* **AC #2** — :meth:`DocxCitationRenderer.render` returns the
  ``(bytes, shadow_md)`` pair the :class:`~eurpe.export.ExportService`
  uses to ship a ``.docx`` file plus run the existing citation audit
  against a textual mirror of the document. The bytes are a valid
  Office Open XML document that opens in Word, Pages, and LibreOffice
  without conversion.
* **AC #3** — citations and source-status labels survive the round
  trip. Each citation row in the rendered ``References`` table carries
  the same ``STATUS_LABEL`` text as the Markdown renderer, and each
  ``Notes`` bullet carries the same ``STATUS_BADGE`` glyph + label.
  Non-funded citations emit the ``STATUS_CAVEAT`` italic line so the
  cautionary framing is visible in the DOCX too.

Why a shadow Markdown string
----------------------------
The export service already runs :class:`~eurpe.generation.CitationAudit`
on rendered Markdown to enforce PRD §22 (release-blocking source-status
preservation). The audit's checks ("every cited STATUS_BADGE appears
in the rendered output", "every ``[N]`` marker in the draft text
appears in the rendered body") only need a string view of the
rendered document — they do not depend on the binary format.

Rather than build a second audit path that opens the DOCX with
python-docx and walks the paragraph tree (fragile against python-docx
version changes), the renderer emits a ``shadow_md`` string alongside
the bytes. ``shadow_md`` is a faithful textual rendering of the same
content the DOCX paragraphs carry, so feeding it through the existing
:meth:`CitationAudit.audit_rendered` exercises the same release-blocking
guarantee. The service stuffs the bytes into
:attr:`ExportResult.content_bytes` and the shadow string into
:attr:`ExportResult.content`; the audit sees the shadow.

Determinism
-----------
The shadow string is byte-equal to ``MarkdownCitationRenderer``'s
output for the same draft. The DOCX bytes are NOT byte-deterministic
across python-docx versions (the package writes ``app.xml`` /
``core.xml`` metadata that includes its own version string), so tests
must not assert on the binary payload's bytes — they assert on the
shadow string and on text round-tripped from the docx via
``docx.Document(BytesIO(bytes))``. Tests live in
``tests/export/test_docx_renderer.py``.
"""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.document import Document as DocxDocument

from eurpe.generation.models import CitationRef, GenerationDraft
from eurpe.generation.render import (
    STATUS_BADGE,
    STATUS_CAVEAT,
    STATUS_LABEL,
    MarkdownCitationRenderer,
)
from eurpe.schema import SectionType


def _section_title(section_type: SectionType) -> str:
    """``METHODOLOGY`` → ``"Methodology"``; ``IMPACT_PATHWAY`` → ``"Impact Pathway"``.

    Mirrors :func:`eurpe.generation.render._section_title` so the DOCX
    heading matches the Markdown heading the user already learnt from
    the CLI output. Kept as a module-private helper rather than imported
    to avoid coupling DOCX rendering to a Markdown helper that might
    diverge in the future.
    """

    return section_type.value.replace("_", " ").title()


def _humanize_programme(programme) -> str:  # type: ignore[no-untyped-def]
    """``HORIZON_EUROPE`` → ``"Horizon Europe"``.

    Same shape as :func:`eurpe.generation.render._humanize_programme`.
    Typed loosely so callers can hand a :class:`~eurpe.schema.Programme`
    member without an extra cast.
    """

    return programme.value.replace("_", " ").title()


def _format_page(page: int | None) -> str:
    """Page placeholder so missing pagination doesn't render as ``None``."""

    return f"p. {page}" if page is not None else "p. n/a"


def _format_section(section: str | None) -> str:
    """Section heading placeholder."""

    return section if section else "n/a"


def _format_title(title: str | None) -> str:
    """Proposal title placeholder."""

    return title if title else "untitled"


class DocxCitationRenderer:
    """Render a :class:`GenerationDraft` to a DOCX byte stream + shadow Markdown.

    Output structure (mirrors :class:`MarkdownCitationRenderer`):

    * Heading 1: section title (e.g., ``"Methodology"``).
    * Body paragraphs: the draft text with inline ``[N]`` markers
      preserved verbatim. Blank lines between paragraphs are honoured
      because the audit's marker-count check sees the same body shape
      the user copy-pastes from Word into the template.
    * Heading 2: ``"References"``.
    * A 7-column table with one row per citation
      (``# / Status / Programme / Call / Section / Page / Source``).
    * Heading 3: ``"Notes"``.
    * One bullet per citation:
      ``[N] {BADGE} — Programme call CALL-ID, p. P, §Section``.
    * Non-funded citations get an italic caveat line beneath the bullet.

    The renderer is stateless and safe to share across calls. The
    accompanying ``shadow_md`` string is produced by reusing
    :class:`MarkdownCitationRenderer` so the audit (which reads
    ``shadow_md`` via :meth:`ExportService.export_section`) sees the
    exact same surface it sees for the Markdown branch.

    Why we reuse :class:`MarkdownCitationRenderer` for the shadow
    -------------------------------------------------------------
    Maintaining two textual renderings of the same draft is a recipe
    for drift: a future tweak to the Markdown table layout would have
    to be mirrored here, and a regression in one path would silently
    pass tests against the other. Delegating to
    ``MarkdownCitationRenderer.render`` keeps the shadow string the
    single source of truth for the audit, while the DOCX paragraph
    tree exists purely for Word/LibreOffice consumption.
    """

    def __init__(
        self,
        *,
        markdown_renderer: MarkdownCitationRenderer | None = None,
    ) -> None:
        # Allow injection so tests can pass a stub renderer when
        # exercising the shadow-string seam in isolation. Defaults to a
        # fresh ``MarkdownCitationRenderer`` because constructing one is
        # cheap and the renderer holds no state.
        self._markdown_renderer = markdown_renderer or MarkdownCitationRenderer()

    def render(self, draft: GenerationDraft) -> tuple[bytes, str]:
        """Render ``draft`` to DOCX bytes plus a shadow Markdown string.

        Returns a ``(bytes, shadow_md)`` tuple so the export service
        can ship the binary payload and run the citation audit on the
        textual mirror in a single round trip.

        Empty citation list is handled identically to the Markdown
        renderer: the ``References`` heading is still written, the
        table is omitted, and a single italic paragraph
        (``"No citations."``) takes its place. This keeps the
        section ordering predictable for downstream consumers and
        keeps the audit (which treats the empty-citation case as
        legal) happy.
        """

        document = Document()
        self._write_heading(document, _section_title(draft.section_type), level=1)
        self._write_body(document, draft.text)
        self._write_heading(document, "References", level=2)
        if draft.citations:
            self._write_references_table(document, draft.citations)
            self._write_heading(document, "Notes", level=3)
            self._write_notes(document, draft.citations)
        else:
            # Italic single line so the reader sees the same signal the
            # Markdown renderer emits (``_No citations._``).
            paragraph = document.add_paragraph()
            run = paragraph.add_run("No citations.")
            run.italic = True

        buffer = BytesIO()
        document.save(buffer)
        shadow_md = self._markdown_renderer.render(draft)
        return buffer.getvalue(), shadow_md

    # ------------------------------------------------------------------
    # internal helpers — small, named, exercised indirectly via render()
    # ------------------------------------------------------------------

    @staticmethod
    def _write_heading(document: DocxDocument, text: str, *, level: int) -> None:
        """Add a heading paragraph at the requested outline level.

        python-docx's ``add_heading`` accepts level 0-9 (0 is the
        document title style). We only use 1-3 here so the section
        title is the document's structural root and References/Notes
        sit below it.
        """

        document.add_heading(text, level=level)

    @staticmethod
    def _write_body(document: DocxDocument, body: str) -> None:
        """Add the draft text as paragraphs, preserving inline ``[N]`` markers.

        We split on double newline so blank-separated paragraphs in
        the draft become distinct paragraphs in the DOCX. Single
        newlines within a paragraph are kept as-is (python-docx
        replaces them with soft line breaks in the saved XML, which is
        the right behaviour for citations that wrap mid-sentence).

        Empty bodies are not expected (Pydantic's ``min_length=1`` on
        ``GenerationDraft.text`` rules them out) but the guard below
        is cheap insurance — an empty string would otherwise create a
        spurious empty paragraph.
        """

        # Strip trailing whitespace so the body doesn't end with an
        # empty paragraph that pushes the References heading down a
        # page. Matches MarkdownCitationRenderer.render which calls
        # ``body.rstrip()`` for the same reason.
        cleaned = body.rstrip()
        if not cleaned:
            return
        for paragraph_text in cleaned.split("\n\n"):
            document.add_paragraph(paragraph_text)

    @staticmethod
    def _write_references_table(
        document: DocxDocument,
        citations: list[CitationRef],
    ) -> None:
        """Write the 7-column references table.

        The column order mirrors the Markdown table exactly so an
        operator who learnt the layout from a CLI render sees the
        same shape here. Header cells are bolded; we set the table
        style to ``"Table Grid"`` so borders render in Word and
        LibreOffice without needing a custom style on the receiving
        side.
        """

        headers = ["#", "Status", "Programme", "Call", "Section", "Page", "Source"]
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        for cell, header in zip(header_cells, headers, strict=True):
            # Clear the default empty paragraph python-docx puts in
            # each new cell and write the header bolded.
            cell.text = ""
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(header)
            run.bold = True

        for citation in citations:
            row = table.add_row().cells
            row[0].text = str(citation.citation_id)
            row[1].text = STATUS_LABEL[citation.source_status]
            row[2].text = _humanize_programme(citation.programme)
            row[3].text = citation.call_id
            row[4].text = _format_section(citation.section_heading)
            # Page cell is the bare number in the table (the "p."
            # prefix is a Notes-block convention — same as the
            # Markdown renderer).
            row[5].text = str(citation.page) if citation.page is not None else "n/a"
            row[6].text = _format_title(citation.proposal_title)

    @staticmethod
    def _write_notes(document: DocxDocument, citations: list[CitationRef]) -> None:
        """Write the per-citation notes block with badges and caveats.

        Each citation produces one ``List Bullet``-styled paragraph
        with the same shape as the Markdown renderer:

            [N] {BADGE} — Programme call CALL-ID, p. P, §Section

        Non-funded citations get an additional italic paragraph
        underneath that carries the ``STATUS_CAVEAT`` text so the
        cautionary framing survives the format conversion.
        """

        for citation in citations:
            badge = STATUS_BADGE[citation.source_status]
            programme = _humanize_programme(citation.programme)
            section = _format_section(citation.section_heading)
            page = _format_page(citation.page)
            line = (
                f"[{citation.citation_id}] {badge} — {programme} call "
                f"{citation.call_id}, {page}, §{section}"
            )
            document.add_paragraph(line, style="List Bullet")

            caveat = STATUS_CAVEAT[citation.source_status]
            if caveat:
                # Strip the surrounding underscores Markdown uses to
                # italicise — python-docx italicises via the run flag.
                # The visible text is identical to the Markdown form
                # so reviewers see the same caveat in both.
                cleaned_caveat = caveat.strip("_")
                paragraph = document.add_paragraph()
                run = paragraph.add_run(cleaned_caveat)
                run.italic = True
